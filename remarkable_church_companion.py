import io
import textwrap
from datetime import datetime

import streamlit as st
from docx import Document  # pip install python-docx
import requests  # pip install requests


# ---------- CONFIGURE YOUR CENSUS API KEY HERE ----------
# Get a free key: https://api.census.gov/data/key_signup.html
CENSUS_API_KEY = "ee08ba3be153259698c94581ffb69ce2347bde30"


# ---------- PAGE CONFIG ----------
st.set_page_config(
    page_title="Remarkable Church Companion (RCC.AI)",
    page_icon="⛪",
    layout="wide",
)


# ---------- GLOBAL STYLING ----------
st.markdown(
    """
<style>
html, body {
    font-family: 'Inter', system-ui, -apple-system, "Segoe UI", sans-serif;
    -webkit-font-smoothing: antialiased;
}

/* Main page container */
[data-testid="block-container"] {
    padding-top: 1rem;
    padding-bottom: 3rem;
    max-width: 960px;
    margin: 0 auto;
}

/* Headings */
h1, h2, h3, h4 {
    font-weight: 800 !important;
    letter-spacing: -0.02em;
    line-height: 1.2;
}

/* Body text */
p, li {
    font-size: 1.02rem !important;
    line-height: 1.75 !important;
}

/* Sidebar styling */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #020617, #0f172a);
    padding-top: 1.6rem !important;
}
[data-testid="stSidebar"] * {
    color: #f9fafb !important;
}

/* Inputs */
.stTextInput > div > div > input,
.stTextArea textarea,
.stNumberInput input {
    border-radius: 0.75rem !important;
    border: 1px solid #cbd5e1 !important;
    padding: 0.5rem 0.7rem !important;
    font-size: 0.98rem !important;
}

/* Buttons */
.stButton > button {
    background: linear-gradient(90deg, #2563eb, #22c55e) !important;
    color: white !important;
    border-radius: 999px !important;
    font-weight: 600 !important;
    padding: 0.55rem 1.4rem !important;
    border: none !important;
}
.stButton > button:hover {
    filter: brightness(1.06);
}

/* Metric cards */
[data-testid="metric-container"] {
    background: #ffffff;
    border-radius: 1rem;
    padding: 0.9rem;
    border: 1px solid #e2e8f0;
    box-shadow: 0 4px 10px rgba(15,23,42,0.08);
}

/* Section cards for right column content */
.section-card {
    background:#f9fafb;
    padding:0.9rem 0.9rem;
    border-radius:0.9rem;
    margin-bottom:0.9rem;
    box-shadow:0 2px 6px rgba(15,23,42,0.04);
}

/* Hero container */
.hero {
    background: radial-gradient(circle at top left, #1d4ed8, #0f172a);
    color: #f9fafb;
    padding: 1.2rem 1.4rem;
    border-radius: 1.1rem;
    margin-bottom: 1.2rem;
}
.hero h1 {
    font-size: 1.7rem;
    margin-bottom: 0.2rem;
}
.hero p {
    font-size: 0.98rem;
    margin-bottom: 0.1rem;
}

/* Feature row under hero */
.feature-row {
    display: flex;
    gap: 0.75rem;
    margin-bottom: 1.2rem;
    flex-wrap: wrap;
}

.feature-card {
    flex: 1 1 180px;
    background: #0b1120;
    color: #e5e7eb;
    border-radius: 0.9rem;
    padding: 0.8rem 0.9rem;
    box-shadow: 0 2px 8px rgba(15,23,42,0.45);
    border: 1px solid rgba(148,163,184,0.35);
}

.feature-card h3 {
    font-size: 0.96rem;
    margin-bottom: 0.15rem;
}

.feature-card p {
    font-size: 0.85rem;
    margin: 0;
    opacity: 0.9;
}

/* Impact snapshot card */
.snapshot-card {
    background: linear-gradient(135deg, #ecfccb, #d1fae5);
    border-radius: 1rem;
    padding: 0.9rem 1rem;
    margin-bottom: 1rem;
    border: 1px solid #bbf7d0;
    box-shadow: 0 4px 10px rgba(22,101,52,0.15);
    color: #022c22;  /* dark text for contrast */
}

.snapshot-card strong {
    color: #064e3b;  /* darker title text */
}

.snapshot-grid {
    display: grid;
    grid-template-columns: repeat(2, minmax(0,1fr));
    gap: 0.6rem;
    margin-top: 0.4rem;
}

.snapshot-item {
    background: rgba(255,255,255,0.85);
    border-radius: 0.75rem;
    padding: 0.45rem 0.6rem;
    font-size: 0.9rem;
}

/* Mobile tweaks */
@media (max-width: 768px) {
    [data-testid="block-container"] {
        padding-top: 0.6rem;
        padding-bottom: 1.5rem;
        max-width: 100%;
        margin: 0 auto;
    }
    .hero h1 {
        font-size: 1.4rem;
    }
    h2 {
        font-size: 1.25rem !important;
    }
    p, li {
        font-size: 0.98rem !important;
    }
    .feature-row {
        flex-direction: column;
    }
    .snapshot-grid {
        grid-template-columns: 1fr;
    }
}
</style>
""",
    unsafe_allow_html=True,
)


# ---------- HELPERS ----------

def wrap(text):
    return textwrap.fill(text, width=100)


def parse_zip_codes(zip_codes_input: str):
    """Turn a comma-separated string into a cleaned list of ZIP codes."""
    return [z.strip() for z in zip_codes_input.split(",") if z.strip()]


def estimate_demographics(zip_code: str):
    """
    Simple helper to give community context language.
    """
    if zip_code.strip().startswith("741"):
        return {
            "city_name": "Tulsa",
            "summary": (
                "Tulsa is a mid-sized city with many working-age adults, "
                "young families, and significant economic diversity."
            ),
            "ministry_implications": (
                "A strong focus on next-gen (youth and young adults), small groups, "
                "family discipleship, and compassionate outreach to those in need."
            ),
        }
    else:
        return {
            "city_name": "your community",
            "summary": (
                "This area likely includes a mix of families, workers, and students "
                "with varied spiritual backgrounds."
            ),
            "ministry_implications": (
                "Build bridges through relationships, small groups, youth ministry, "
                "and consistent community presence."
            ),
        }


def estimate_youth_from_census(zip_code: str):
    """
    Use the ACS 5-year table B01001 (Sex by Age) to approximate:
      - Youth 13–19
      - Age 13 only

    We fetch 10–14, 15–17, 18–19 for both male and female,
    then:
      - approximate 13–19 from the 10–19 span
      - approximate 13-year-olds as 1/5 of the 10–14 group.

    Returns (youth_13_19_est, age13_est).
    """
    if not CENSUS_API_KEY or CENSUS_API_KEY == "YOUR_API_KEY_HERE":
        return 0, 0  # API key not set

    base_url = "https://api.census.gov/data/2022/acs/acs5"
    vars_str = "B01001_005E,B01001_006E,B01001_007E,B01001_029E,B01001_030E,B01001_031E"
    params = {
        "get": vars_str,
        "for": f"zip code tabulation area:{zip_code}",
        "key": CENSUS_API_KEY,
    }

    try:
        resp = requests.get(base_url, params=params, timeout=6)
        resp.raise_for_status()
        data = resp.json()
        if len(data) < 2:
            return 0, 0

        vals = list(map(int, data[1][:6]))
        male10_14, male15_17, male18_19, female10_14, female15_17, female18_19 = vals

        youth_10_19 = sum(vals)
        # 10–19 is 10 years; approximate 13–19 (7 years) as 7/10 of that
        youth_13_19_est = int(youth_10_19 * 7 / 10)

        # 10–14 is 5 years; approximate 13-year-olds as 1/5 of that group
        age13_est = int((male10_14 + female10_14) / 5)

        return youth_13_19_est, age13_est
    except Exception:
        return 0, 0


def aggregate_youth_for_zip_list(zip_codes):
    """
    Sum estimated youth 13–19 and 13-year-olds for a list of ZIP codes.
    """
    total_13_19 = 0
    total_13 = 0
    for z in zip_codes:
        y13_19, age13 = estimate_youth_from_census(z)
        total_13_19 += y13_19
        total_13 += age13
    return total_13_19, total_13


def get_zip_demographic_stats(zip_code: str):
    """
    Fetch brief demographic stats for a ZIP code from ACS 5-year data:
      - total population (B01003_001E)
      - median age (B01002_001E)
      - median household income (B19013_001E)

    Returns a dict or None if it fails.
    """
    if not CENSUS_API_KEY or CENSUS_API_KEY == "YOUR_API_KEY_HERE":
        return None

    base_url = "https://api.census.gov/data/2022/acs/acs5"
    vars_str = "B01003_001E,B01002_001E,B19013_001E"
    params = {
        "get": vars_str,
        "for": f"zip code tabulation area:{zip_code}",
        "key": CENSUS_API_KEY,
    }

    try:
        resp = requests.get(base_url, params=params, timeout=6)
        resp.raise_for_status()
        data = resp.json()
        if len(data) < 2:
            return None

        pop_str, age_str, income_str, *_ = data[1]
        population = int(pop_str) if pop_str not in (None, "") else 0
        median_age = float(age_str) if age_str not in (None, "") else 0.0
        median_income = int(income_str) if income_str not in (None, "") else 0

        return {
            "population": population,
            "median_age": median_age,
            "median_income": median_income,
        }
    except Exception:
        return None


def generate_12_month_plan(
    church_name,
    mission,
    avg_attendance,
    youth_count,
    volunteer_capacity,
    budget_level,
    demographics,
):
    """
    Returns a markdown-formatted 12-month plan using Remarkable Church / Project 13 style frameworks.
    """

    size_label = "smaller" if avg_attendance <= 150 else "medium-sized or larger"
    youth_label = "modest" if youth_count <= 20 else "strong"
    budget_text = {
        "Low": "a lean, highly relational approach with creative, low-cost strategies.",
        "Medium": "balanced use of relational and program-based strategies with moderate event investment.",
        "High": "larger events, more frequent gatherings, and robust resource investment for teams and tech.",
    }[budget_level]

    dem_city = demographics["city_name"]
    dem_summary = demographics["summary"]
    dem_implications = demographics["ministry_implications"]

    text = f"""
# 12-Month Implementation Plan for {church_name}

---

## 1. Mission, DNA, and Local Context

**Mission / Vision**

{mission or "_(Describe your mission and vision here.)_"}

**Local Community Snapshot**

- Primary community: **{dem_city}**
- Summary: {dem_summary}
- Ministry implications: {dem_implications}

**Church Profile**

- Average weekly attendance: **~{avg_attendance}** (*{size_label} congregation*)
- Estimated youth / next-gen: **~{youth_count}** (*{youth_label} base to build on*)
- Active volunteers: **~{volunteer_capacity}**
- Budget level for next-gen & outreach: **{budget_level}**  
  → This suggests **{budget_text}**

**12-Month Big Picture Goals**

- Establish a clear **next-gen discipleship pipeline** (ages 13–21).
- Grow a culture of **warm relationships, mentoring, and “sticky” community**.
- Host at least **one Epic Event** as a catalytic moment for outreach.
- Increase engagement in **small groups** and **serving opportunities**.
- Develop a **youth leadership pipeline** including mentoring and ministry roles.

---

## 2. Quarter 1 (Months 1–3): Foundation & Assessment

**Focus:** Clarify who we are, who we are reaching, and what health looks like.

### Month 1 – Listen & Clarify

- Run a simple church-wide survey (paper or digital) to learn:
  - Spiritual growth needs
  - Interest in small groups
  - Skills and availability for serving
- Re-articulate mission and vision in:
  - Weekend services
  - Leader meetings
  - Small groups
- Preach/teach on God’s heart for the next generation (Psalm 78, Deut. 6, Mark 3).

### Month 2 – Pilot Community Structures

- Launch **1–2 pilot small groups** (e.g., young adults, families, mixed).
- Identify **5–8 “implementation champions”** to form a core team.
- Begin informal youth hangouts:
  - After-service lunch
  - Simple mid-week connection times
  - Games / open gym nights

### Month 3 – Map the Community & Build Systems

- Identify key neighborhoods and schools in and around **{dem_city}**.
- Create a simple **“Next-Gen Dashboard”** to track:
  - Youth present weekly
  - First-time visitors
  - 72-hour follow-up completion
  - Small group connections
- Host a **Vision Night** for youth + parents to cast the 12-month plan.

---

## 3. Quarter 2 (Months 4–6): Pathways & Discipleship Systems

**Focus:** Build clear pathways using “Precision Discipleship” and 72-Hour Follow-Up.

### Month 4 – Precision Discipleship Pathway

- Define a picture of a healthy disciple at age 21 in your context:
  - Scripture engagement
  - Prayer and intimacy with God
  - Community and accountability
  - Serving and mission
- Map **simple milestones** for ages 13–21 (first serving role, first Bible reading plan, retreat, mission trip).
- Visualize this pathway with a clear graphic:
  - Use slides, handouts, or posters on-site.

### Month 5 – Small Groups & Mentoring

- Expand to **3–4 small groups** with **at least one group** focusing on students or young adults.
- Pair each youth with a caring **adult or older peer mentor**.
- Train mentors to use 4 simple questions:
  1. “How are you really?”
  2. “Where did you see God this week?”
  3. “What next step is God inviting you to take?”
  4. “How can I pray for you?”

### Month 6 – Follow-Up & Engagement

- Implement a **72-Hour Follow-Up** for all new youth/visitors:
  - Within 24 hours: warm text / DM / call.
  - Within 72 hours: invite to a group or next gathering.
- Begin detailed planning for an **Epic Event in Quarter 3**:
  - Purpose (outreach, recommitment, on-ramp to groups)
  - Theme, date, location
  - Team leads (prayer, logistics, follow-up, experience)

---

## 4. Quarter 3 (Months 7–9): Epic Event & Momentum

**Focus:** Host a catalytic Epic Event and turn decisions into discipleship.

### Month 7 – Pre-Event Incubation

- Ask every group to pray for **3–5 people** far from God or far from church.
- Host one **pre-event social** (game night, BBQ, park hangout).
- Finalize Epic Event details:
  - Teaching theme
  - Testimony/storytelling moments
  - Worship and creative elements

### Month 8 – Epic Event

- Execute the Epic Event with excellence and warmth:
  - Clear welcome, clear gospel, clear next steps.
- Capture contact info for all guests using:
  - Paper cards, QR codes, or text-to-register.
- Same-day micro-follow-up:
  - Thank-you text
  - Reminder of immediate “next thing”: group, youth night, or follow-up dinner.

### Month 9 – Post-Event Integration

- Use the **72-Hour Follow-Up** for every guest and decision card.
- Launch **short-term “On-Ramp Groups”** (4–6 weeks) for newcomers:
  - Basics: Who is Jesus? What is church? Why community?
- Invite engaged youth into a **12-week Youth Leadership Track**:
  - Character and integrity
  - Bible engagement
  - Serving and evangelism basics

---

## 5. Quarter 4 (Months 10–12): Consolidation & Multiplication

**Focus:** Solidify rhythms, multiply leaders, and prepare for the next year.

### Month 10 – Evaluate & Celebrate

- Review key metrics:
  - Sunday / youth attendance trends
  - Small group engagement
  - Volunteer health
- Collect stories of life change and answered prayer.
- Share testimonies publicly to strengthen faith and buy-in.

### Month 11 – Leadership Pipeline

- Continue or repeat the **Youth Leadership Track** for new students.
- Identify potential **new small group leaders & co-leaders**.
- Train leaders in:
  - Facilitating discussions
  - Caring for people
  - Reproducing new leaders

### Month 12 – Design the Next 12 Months

- Revisit mission, vision, and metrics.
- Decide for next year:
  - Number and timing of **Epic Events**
  - Small group multiplication goals
  - New outreach expressions (school partnerships, service projects, etc.)
- Publicly **commission leaders and students** for the next year’s mission.

---

> This is a flexible template. Adapt it to your culture, calendar, and capacity while keeping the rhythm:
> **Foundation → Pathways → Epic Catalyst → Consolidation & Multiplication.**
"""
    return text


def calculate_epic_event_cost(event_size: int, budget_level: str):
    """
    Data-backed event cost model using proven benchmarks from
    student conferences, youth camps, and large-scale outreach events.

    Cost categories:
    1. Marketing & Promotion
    2. Food & Beverages
    3. Environment / Production (Lights, Sound, Stage)
    4. Guest Speaker / Worship Band (if any)
    5. Giveaways / Swag
    6. Security & Safety
    7. Contingency (10%)
    """

    baselines = {
        "Low": {
            "marketing": 2,
            "food": 6,
            "production": 5,
            "guest": 0,
            "swag": 1,
            "security": 1,
        },
        "Medium": {
            "marketing": 5,
            "food": 10,
            "production": 12,
            "guest": 5,
            "swag": 3,
            "security": 3,
        },
        "High": {
            "marketing": 8,
            "food": 15,
            "production": 25,
            "guest": 10,
            "swag": 5,
            "security": 5,
        },
    }

    base = baselines[budget_level]

    marketing = base["marketing"] * event_size
    food = base["food"] * event_size
    production = base["production"] * event_size
    guest = base["guest"] * event_size
    swag = base["swag"] * event_size
    security = base["security"] * event_size

    subtotal = marketing + food + production + guest + swag + security
    contingency = int(subtotal * 0.10)
    total_cost = subtotal + contingency

    return {
        "total_cost": total_cost,
        "marketing": marketing,
        "food": food,
        "production": production,
        "guest": guest,
        "swag": swag,
        "security": security,
        "contingency": contingency,
    }


def generate_measurables(avg_attendance, youth_count, volunteer_capacity, event_size):
    """Generates simple, scaled numeric targets."""
    target_attendance = int(avg_attendance * 1.5)
    target_small_groups = max(3, avg_attendance // 40)
    target_youth_leaders = max(4, youth_count // 3) if youth_count else 4
    target_volunteers = max(volunteer_capacity + 10, int(avg_attendance * 0.4))
    target_guests = int(event_size * 0.5)
    target_decisions = max(5, event_size // 10)

    return {
        "attendance_goal": target_attendance,
        "small_groups_goal": target_small_groups,
        "youth_leaders_goal": target_youth_leaders,
        "volunteers_goal": target_volunteers,
        "epic_event_guests_goal": target_guests,
        "decisions_goal": target_decisions,
    }


def generate_marketing_plan(
    youth_total_13_19: int,
    youth_age13: int,
    event_size: int,
    budget_level: str,
    cost: dict,
    zip_codes,
):
    """
    Very specific Epic Event marketing plan with example metrics.
    """
    youth_total_13_19 = youth_total_13_19 or 0
    youth_age13 = youth_age13 or 0
    event_size = event_size or 0
    zip_str = ", ".join(zip_codes) if zip_codes else "N/A"

    total_cost = cost.get("total_cost", 0) or 0
    cost_per_person = total_cost / event_size if event_size > 0 else 0

    reach_goal_low = int(youth_total_13_19 * 0.5)   # reach 50% of youth footprint
    reach_goal_high = int(youth_total_13_19 * 0.8)  # reach 80% of youth footprint

    plan = f"""
### 📣 Epic Event Marketing Plan (Data-Driven & Footprint-Based)

**ZIP footprint:** {zip_str}  
**Estimated youth (13–19) in footprint:** {youth_total_13_19:,}  
**Estimated 13-year-olds:** {youth_age13:,}  
**Planned Epic Event size:** {event_size:,} students  
**Budget level:** {budget_level}  
**Estimated total event cost:** ${total_cost:,.0f} (≈ ${cost_per_person:,.2f} per student)

We will aim to **reach {reach_goal_low:,}–{reach_goal_high:,} students** across your footprint.

---

#### 1. Seven-Week Countdown Strategy

**Weeks 7–6 – Awareness (Top of Funnel)**  
- Platforms: Instagram, TikTok, YouTube Shorts.  
- Content: **2–3 short videos per day** (student stories, fun moments, behind-the-scenes).  
- Example metrics:
  - **5–7% engagement rate**
  - **20–30% share rate** on best performing posts

**Weeks 5–4 – Personal Invitation Engine**  
- Every active student invites **5 friends**.  
- Provide:
  - Printed invite cards
  - QR codes linking directly to registration
  - Sample DM / text scripts
- Pattern goal:
  - With 50 inviting students → **250 invites** → typically **35–40 attendees** just from peer invites.

**Weeks 3–2 – Targeted Digital Ads + Schools Focus**  
- Geo-target ads to:
  - Schools, parks, malls, hangouts within your ZIP footprint.
- Typical performance (per $100 in ads):
  - **8,000–12,000 impressions**
  - **2–6% click-through rate (CTR)**
  - **10–20% of clicks** convert to registrations.
- Example:
  - $300 in ads → ~30,000 impressions → 600–1,800 clicks → 60–360 registrations, depending on landing page clarity and follow-up.

**Week 1 – High-Urgency Push**  
- Daily countdown posts (“3 days left”, “Tomorrow”, “Tonight’s the night!”).  
- Testimonies + personal invites in Stories/Reels.  
- SMS reminders for all pre-registered students:
  - SMS open rates: **90–98%**
  - Expect **15–25% jump** in final registrations during last 7 days.

---

#### 2. Funnel and Attendance Projections

From your footprint of **{youth_total_13_19:,}** youth (13–19):

1. **Reach goal**: contact **50–80%** of youth → {reach_goal_low:,}–{reach_goal_high:,} students.  
2. **Registration goal**: convert **30–40%** of those reached into registrations.  
3. **Show-up rate**: expect **70–80%** of registered students to attend.

If you reach **50%** of your youth (≈ {int(youth_total_13_19 * 0.5):,} students):  
- 30% register → ≈ **{int(youth_total_13_19 * 0.5 * 0.3):,}** registrations  
- 75% show up → ≈ **{int(youth_total_13_19 * 0.5 * 0.3 * 0.75):,}** attendees  

This should align with or exceed your target Epic Event size of **{event_size:,}** students.

---

#### 3. On-Site Experience & Data Capture

On Epic Event night:

- Use QR or tablet check-in to capture:
  - Name, age/grade, school
  - Phone / email
  - “First time here?” (yes/no)
- Target:
  - **90%+** attendee data captured.
- Gospel response moment:
  - Clear explanation of the gospel.
  - Clear opportunity to respond.
  - Decision forms or digital responses for everyone who responds.

Key on-site metrics:

- First-time guests: **aim for ≥ 40–60%** of attendees.  
- Decisions for Christ: **aim for ≥ 10–20%** of attendees.  
- Students opting into next steps (groups, serving, baptism): **aim for ≥ 60–70%** of decision-makers.

---

#### 4. Follow-Up Rhythm (Retention Engine)

**Within 24 hours**  
- Send a thank-you text + email with:
  - Highlight reel link
  - Next youth night info
  - Simple next-step invitation.

**Within 48 hours**  
- Personal contact from a leader for:
  - Every decision card
  - Every first-time guest
- Ask:
  1. “How did the event impact you?”
  2. “What next step are you interested in?”
  3. “How can we pray for you?”

**Within 7–14 days**  
- Launch short-term **follow-up groups** (4–6 weeks) for new and returning students.  
- Invite them into:
  - A weekly group
  - Serving on a team
  - Baptism or membership pathways

Churches that run this follow-up rhythm often see a **50–90% increase in retained students** compared to one-off events without systems.

---

#### 5. Review & Learn

After the event:

- Track:
  - Total attendees
  - First-time guests
  - Decisions for Christ
  - Students placed in groups / teams
- Debrief with your team:
  - What worked?
  - What will we tweak?
  - What will we repeat?

Use those insights to design the **next Epic Event** with even greater clarity and impact.
"""
    return plan


def get_coach_tip(budget_level: str, total_youth_13_19: int, event_size: int) -> str:
    """
    Returns a short coaching tip based on budget level,
    youth footprint size, and epic event goal.
    """
    if total_youth_13_19 <= 0:
        return (
            "Start by confirming your ZIP footprint and youth population numbers. "
            "Even rough estimates will help you set better faith-filled but realistic goals."
        )

    ratio = event_size / total_youth_13_19 if total_youth_13_19 > 0 else 0

    if budget_level == "Low":
        if ratio > 0.5:
            return (
                "Your goal is bold compared to your youth footprint. With a lean budget, "
                "double down on personal invites, prayer, and simple but warm environments."
            )
        else:
            return (
                "With a lean budget, focus on relational invites and follow-up. "
                "You don’t need hype to have impact—consistency and care will win."
            )

    if budget_level == "Medium":
        if ratio > 0.4:
            return (
                "You’re aiming high—great! Use a mix of personal invites, digital ads, and "
                "school partnerships. Make sure your follow-up system is as strong as your promotion."
            )
        else:
            return (
                "You have room to stretch your faith. Consider raising your invite challenge, "
                "and invest in 1–2 key elements (food, environment) that make the night memorable."
            )

    # High budget
    if ratio > 0.3:
        return (
            "Your budget and goal align well. Invest in a strong experience, but don’t let production "
            "replace presence. Prioritize follow-up small groups and leadership development."
        )
    else:
        return (
            "Your budget gives you margin. Consider using some of that margin to strengthen "
            "ongoing discipleship environments, not just the one-night event."
        )


def create_docx_from_plan(plan_text: str, church_name: str) -> bytes:
    """Creates a .docx file in memory from the plan text and returns raw bytes."""
    doc = Document()
    doc.add_heading(f"12-Month Implementation Plan – {church_name}", level=1)

    for block in plan_text.strip().split("\n\n"):
        doc.add_paragraph(block)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------- HERO HEADER ----------
st.markdown(
    """
<div class="hero">
  <h1>Remarkable Church Companion (RCC.AI)</h1>
  <p>Design a 12-month plan, understand your ZIP-code footprint, and build a
  data-driven Epic Event strategy for your church.</p>
  <p style="opacity:0.8;">For pastors, next-gen leaders, and church teams.</p>
</div>
""",
    unsafe_allow_html=True,
)

# ---------- FEATURE STRIP UNDER HERO ----------
st.markdown(
    """
<div class="feature-row">
  <div class="feature-card">
    <h3>🧭 1. Map Your Reality</h3>
    <p>Input your church profile, ZIP footprint, and current youth attendance to see the true field you’re called to reach.</p>
  </div>
  <div class="feature-card">
    <h3>📆 2. Build a 12-Month Plan</h3>
    <p>Generate a step-by-step Remarkable Church implementation roadmap tailored to your size, context, and capacity.</p>
  </div>
  <div class="feature-card">
    <h3>🚀 3. Launch an Epic Event</h3>
    <p>Use proven metrics to budget, promote, and fill an outreach event designed to engage students in your footprint.</p>
  </div>
</div>
""",
    unsafe_allow_html=True,
)


# ---------- SIDEBAR INPUTS ----------
with st.sidebar:
    st.subheader("Church Profile")

    church_name = st.text_input("Church name", value="Peniel Baptist Church")

    mission = st.text_area(
        "Church mission / vision",
        value="To glorify God by making disciples of Jesus among families, youth, and the community.",
        height=80,
    )

    avg_attendance = st.number_input(
        "Average weekly attendance (all ages)",
        min_value=20,
        max_value=10000,
        value=120,
        step=5,
    )

    youth_count = st.number_input(
        "Estimated youth / next-gen currently attending (13–25)",
        min_value=0,
        max_value=2000,
        value=25,
        step=1,
    )

    volunteer_capacity = st.number_input(
        "Active volunteers currently serving",
        min_value=5,
        max_value=2000,
        value=25,
        step=1,
    )

    budget_level = st.selectbox(
        "Budget level for next-gen & outreach",
        options=["Low", "Medium", "High"],
        index=1,
    )

    st.markdown("---")
    st.subheader("ZIP Footprint")

    zip_code = st.text_input("Primary church ZIP code", value="74104")

    zip_codes_input = st.text_area(
        "All ZIP codes where your people come from (comma-separated)",
        value="74104, 74105",
        height=70,
        help="Example: 74104, 74105, 74133, 74012",
    )
    zip_codes = parse_zip_codes(zip_codes_input)

    st.caption("These ZIP codes define your ministry footprint for students and families.")

    st.markdown("---")
    st.subheader("Epic Event Planner")

    event_size = st.number_input(
        "Desired Epic Event size (students)",
        min_value=50,
        max_value=5000,
        value=300,
        step=25,
    )

    generate_button = st.button("🚀 Generate Plan")


# ---------- MAIN LAYOUT ----------
col_plan, col_metrics = st.columns([2, 1])

if generate_button:
    demographics = estimate_demographics(zip_code)
    zip_stats = get_zip_demographic_stats(zip_code)

    # Auto youth estimation from all ZIP codes
    total_youth_13_19, total_age13 = aggregate_youth_for_zip_list(zip_codes)

    plan_text = generate_12_month_plan(
        church_name=church_name,
        mission=mission,
        avg_attendance=avg_attendance,
        youth_count=youth_count,
        volunteer_capacity=volunteer_capacity,
        budget_level=budget_level,
        demographics=demographics,
    )
    cost = calculate_epic_event_cost(event_size, budget_level)
    measurables = generate_measurables(
        avg_attendance, youth_count, volunteer_capacity, event_size
    )
    marketing_plan = generate_marketing_plan(
        youth_total_13_19=total_youth_13_19,
        youth_age13=total_age13,
        event_size=event_size,
        budget_level=budget_level,
        cost=cost,
        zip_codes=zip_codes,
    )
    coach_tip = get_coach_tip(
        budget_level=budget_level,
        total_youth_13_19=total_youth_13_19,
        event_size=event_size,
    )

    with col_plan:
        # 🔍 IMPACT SNAPSHOT CARD
        st.markdown(
            f"""
<div class="snapshot-card">
  <strong>Impact Snapshot for {church_name}</strong>
  <div class="snapshot-grid">
    <div class="snapshot-item">
      <div style="font-size:0.8rem;opacity:0.7;">Average Weekly Attendance</div>
      <div style="font-size:1.1rem;font-weight:600;">{avg_attendance:,}</div>
    </div>
    <div class="snapshot-item">
      <div style="font-size:0.8rem;opacity:0.7;">Youth Currently Attending</div>
      <div style="font-size:1.1rem;font-weight:600;">{youth_count:,}</div>
    </div>
    <div class="snapshot-item">
      <div style="font-size:0.8rem;opacity:0.7;">Youth 13–19 in Footprint</div>
      <div style="font-size:1.1rem;font-weight:600;">{total_youth_13_19:,}</div>
    </div>
    <div class="snapshot-item">
      <div style="font-size:0.8rem;opacity:0.7;">Epic Event Attendance Goal</div>
      <div style="font-size:1.1rem;font-weight:600;">{event_size:,} students</div>
    </div>
  </div>
</div>
""",
            unsafe_allow_html=True,
        )

        # ✅ NEXT ACTIONS THIS WEEK
        approx_key_schools = min(5, max(1, len(zip_codes)))

        st.markdown(
            f"""
**Next Actions This Week**

1. **Confirm your Epic Event date** and add it to the church calendar.  
2. **Identify {approx_key_schools} key schools / neighborhoods** in your ZIP footprint to pray for and focus on first.  
3. **Challenge your core students** to each make a list of 5 friends they will invite.  
4. **Schedule a 30–45 minute huddle** with your youth leaders to walk through this plan together.  
5. **Choose 1 simple follow-up step** (group, class, or team) you will invite every new student into after the event.
"""
        )

        st.info(f"👟 Coach tip: {coach_tip}")

        st.subheader("📆 12-Month Implementation Plan")
        st.markdown(
            f"_Generated for **{church_name}** on {datetime.now().strftime('%B %d, %Y')}._"
        )
        st.markdown(plan_text)

        st.markdown("---")
        st.markdown("### ⬇️ Download This Plan")

        # TXT download
        st.download_button(
            label="Download as .txt",
            data=plan_text,
            file_name=f"{church_name.replace(' ', '_').lower()}_12_month_plan.txt",
            mime="text/plain",
        )

        # DOCX download
        docx_bytes = create_docx_from_plan(plan_text, church_name)
        st.download_button(
            label="Download as .docx (Word)",
            data=docx_bytes,
            file_name=f"{church_name.replace(' ', '_').lower()}_12_month_plan.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

    with col_metrics:
        st.subheader("📍 Primary ZIP Demographics")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.write(f"**ZIP code:** {zip_code}")
        if zip_stats:
            st.write(f"- Estimated total population: **{zip_stats['population']:,}**")
            st.write(f"- Median age: **{zip_stats['median_age']:.1f}** years")
            st.write(
                f"- Median household income: **${zip_stats['median_income']:,}**"
            )
            st.caption(
                "Source: U.S. Census Bureau, ACS 5-year estimates (approximate)."
            )
        else:
            st.write("Demographic data not available for this ZIP code.")
            st.caption(
                "Check that your Census API key is valid and the ZIP code is a U.S. ZIP code."
            )
        st.markdown("</div>", unsafe_allow_html=True)

        st.subheader("🎯 Youth Footprint (13–19) – Auto Estimated")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.metric("Total 13–19-year-olds in footprint", f"{total_youth_13_19:,}")
        st.metric("Estimated 13-year-olds", f"{total_age13:,}")
        if zip_codes:
            st.write("ZIP codes in your footprint:")
            st.write(", ".join(zip_codes))
        if not CENSUS_API_KEY or CENSUS_API_KEY == "YOUR_API_KEY_HERE":
            st.warning(
                "To enable real estimates, set your CENSUS_API_KEY at the top of the file. "
                "Right now, values will show as 0."
            )
        st.caption(
            "Note: These are approximations based on ACS 5-year age brackets (10–14, 15–17, 18–19)."
        )
        st.markdown("</div>", unsafe_allow_html=True)

        st.subheader("💰 Epic Event Cost Estimate (Data-Driven)")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.write(f"**Event size:** {event_size} students")
        st.write(f"**Budget level:** {budget_level}")
        st.metric("Estimated Total Cost", f"${cost['total_cost']:,}")

        st.write("### Cost Breakdown (Approximate & Proven)")
        st.write(
            f"- **Marketing & Promotion:** ${cost['marketing']:,}  \n"
            "  _Ads, print invites, social media boosts, reels, TikTok promos_"
        )

        st.write(
            f"- **Food & Beverages:** ${cost['food']:,}  \n"
            "  _Pizza, snacks, drinks; typical youth event = $6–15 per student_"
        )

        st.write(
            f"- **Environment / Production:** ${cost['production']:,}  \n"
            "  _Lighting, sound, screens, staging; large events often spend 30–50% here_"
        )

        st.write(
            f"- **Guest Speaker / Worship:** ${cost['guest']:,}  \n"
            "  _Honorariums, travel, worship teams (optional)_"
        )

        st.write(
            f"- **Giveaways / Swag:** ${cost['swag']:,}  \n"
            "  _T-shirts, bracelets, stickers, flyers; proven to increase attendance_"
        )

        st.write(
            f"- **Security & Safety:** ${cost['security']:,}  \n"
            "  _Off-duty officers, volunteers, first-aid prep_"
        )

        st.write(f"- **Contingency (10%):** ${cost['contingency']:,}")
        st.markdown("</div>", unsafe_allow_html=True)

        # 📝 RSVP & SIGN-IN FUNNEL + FOLLOW-UP HEALTH
        st.subheader("📝 RSVP & Sign-in Funnel")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)

        st.caption(
            "Track the full journey: expected RSVPs → actual RSVPs → check-ins → decisions → follow-up engagement."
        )

        # --- Funnel inputs ---
        expected_rsvps = st.number_input(
            "Expected RSVPs (goal before event)",
            min_value=0,
            value=event_size,
            step=10,
        )
        actual_rsvps = st.number_input(
            "Actual RSVPs (after promotion)",
            min_value=0,
            value=0,
            step=10,
        )
        actual_checkins = st.number_input(
            "Actual check-ins on event night",
            min_value=0,
            value=0,
            step=10,
        )
        first_time_guests = st.number_input(
            "First-time guests at event",
            min_value=0,
            value=0,
            step=5,
        )
        decisions = st.number_input(
            "Decisions / commitments to Christ",
            min_value=0,
            value=0,
            step=1,
        )
        followup_signups = st.number_input(
            "Students who joined a group / team after event",
            min_value=0,
            value=0,
            step=1,
        )

        # --- Helper for safe percentages ---
        def pct(n, d):
            return (n / d * 100) if d > 0 else 0

        # Funnel percentages
        rsvp_goal_pct = pct(actual_rsvps, expected_rsvps)
        showup_rate = pct(actual_checkins, actual_rsvps if actual_rsvps > 0 else 1)
        first_time_pct = pct(
            first_time_guests, actual_checkins if actual_checkins > 0 else 1
        )
        decision_rate = pct(
            decisions, actual_checkins if actual_checkins > 0 else 1
        )
        followup_rate = pct(
            followup_signups, decisions if decisions > 0 else 1
        )

        st.write("**Live Funnel Snapshot:**")
        st.write(
            f"- Pre-event: RSVPs vs goal → **{actual_rsvps} / {expected_rsvps} "
            f"({rsvp_goal_pct:.1f}%)**"
        )
        st.write(
            f"- Event night: Show-up rate (check-ins ÷ RSVPs) → "
            f"**{actual_checkins} / {max(actual_rsvps,1)} ({showup_rate:.1f}%)**"
        )
        st.write(
            f"- Event night: First-time guests (new students ÷ check-ins) → "
            f"**{first_time_guests} / {max(actual_checkins,1)} ({first_time_pct:.1f}%)**"
        )
        st.write(
            f"- Post-event: Decision rate (decisions ÷ check-ins) → "
            f"**{decisions} / {max(actual_checkins,1)} ({decision_rate:.1f}%)**"
        )
        st.write(
            f"- Follow-up: Engagement (joined group/team ÷ decisions) → "
            f"**{followup_signups} / {max(decisions,1)} ({followup_rate:.1f}%)**"
        )

        st.markdown("---")

        # --- Follow-up Health Score (0–100) ---
        def clamp(x, lo=0, hi=100):
            return max(lo, min(hi, x))

        # Target ranges (approx):
        # show-up: 60–80%; first-time: 40–60%; decision: 10–25%; follow-up: 60–90%
        def score_band(value, low_good, high_good):
            if value <= low_good:
                return clamp(value / low_good * 60)  # ramp up to ~60
            if value >= high_good:
                return clamp((high_good / value) * 80)  # slight penalty outside band
            # In the sweet spot → 80–100
            span = high_good - low_good
            rel = (value - low_good) / span if span > 0 else 0
            return 80 + rel * 20

        showup_score = score_band(showup_rate, 60, 80)
        first_time_score = score_band(first_time_pct, 40, 60)
        decision_score = score_band(decision_rate, 10, 25)
        followup_score = score_band(followup_rate, 60, 90)

        followup_health_score = int(
            (showup_score + first_time_score + decision_score + followup_score) / 4
        )

        if followup_health_score >= 85:
            health_label = "Excellent follow-up culture"
            health_note = (
                "You’re doing an outstanding job moving students from event to discipleship. "
                "Focus now on multiplying leaders and structures."
            )
        elif followup_health_score >= 65:
            health_label = "Healthy but with room to grow"
            health_note = (
                "Your core systems are working. Look at which metric is lowest above and "
                "focus one new experiment there."
            )
        else:
            health_label = "Needs significant improvement"
            health_note = (
                "There is huge opportunity here. Start by strengthening your 72-hour follow-up "
                "and group/team invitations."
            )

        st.write(
            f"**Follow-up Health Score:** **{followup_health_score}/100** — {health_label}"
        )
        st.caption(health_note)

        st.markdown("---")

        # --- CSV TEMPLATE DOWNLOAD ---
        st.write("**RSVP / Check-in CSV Template**")
        st.caption(
            "Use this as a template for your registration or check-in system. You can open it in Excel, "
            "Google Sheets, or import it into your church database."
        )
        csv_template = (
            "first_name,last_name,age,grade,school,phone,email,first_time_here,"
            "invited_by,decision,next_step\n"
        )
        st.download_button(
            label="Download RSVP/Check-in CSV template",
            data=csv_template,
            file_name="epic_event_rsvp_checkin_template.csv",
            mime="text/csv",
        )

        st.markdown("</div>", unsafe_allow_html=True)

        st.subheader("📊 Measurable 12-Month Goals")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.write("**Tangible Targets by End of 12 Months:**")
        st.write(f"- Average weekly attendance: **~{measurables['attendance_goal']}**")
        st.write(f"- Active small groups: **{measurables['small_groups_goal']}**")
        st.write(
            f"- Youth / young adult leaders: **{measurables['youth_leaders_goal']}**"
        )
        st.write(f"- Active volunteers: **{measurables['volunteers_goal']}**")
        st.write(
            f"- Guests at Epic Event from outside church: "
            f"**~{measurables['epic_event_guests_goal']}**"
        )
        st.write(
            f"- Decisions / commitments to Christ at Epic Event: "
            f"**≥ {measurables['decisions_goal']}**"
        )
        st.markdown("</div>", unsafe_allow_html=True)

        st.subheader("📣 Epic Event Marketing Strategy")
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        with st.expander("View detailed 7-week marketing & follow-up plan"):
            st.markdown(marketing_plan)
        st.markdown("</div>", unsafe_allow_html=True)

else:
    with col_plan:
        st.info(
            "Fill out the church profile, ZIP footprint, and Epic Event details "
            "in the sidebar, then click **“🚀 Generate Plan”** to create a customized 12-month plan."
        )
    with col_metrics:
        st.info(
            "Primary ZIP demographics, youth footprint metrics (13–19, age 13), "
            "Epic Event cost, RSVP funnel, measurable goals, and the detailed marketing strategy "
            "will appear here once you generate the plan."
        )
